from docx import Document

# Create a new Document
doc = Document()

# Add Title
doc.add_heading('Beginner to Advanced DevOps Roadmap', level=1)

# Add Sections
sections = [
    ("Version Control Systems (VCS)", [
        ("Learn Git", [
            "Topics: Basics, Branching, Merging, Pull Requests, Conflict Resolution.",
            "Resources:",
            "Official Git Documentation: https://git-scm.com/doc",
            "YouTube: \"Git Tutorial for Beginners by FreeCodeCamp\" (https://www.youtube.com/watch?v=RGOj5yH7evk)"
        ])
    ]),
    ("Build Automation & CI/CD Pipelines", [
        ("Learn Jenkins", [
            "Topics: Jobs, Pipelines, Plugins, Declarative vs. Scripted Pipelines.",
            "Resources:",
            "Official Jenkins Tutorials: https://www.jenkins.io/doc/tutorials/",
            "YouTube: \"Jenkins Tutorial for Beginners by Edureka\" (https://www.youtube.com/watch?v=FX322RVNGj4)"
        ]),
        ("Explore GitHub Actions", [
            "Topics: Creating Workflows, Secrets, Automation.",
            "Resources:",
            "YouTube: \"GitHub Actions Crash Course by Traversy Media\" (https://www.youtube.com/watch?v=R8_veQiYBjI)"
        ])
    ]),
    ("Containerization", [
        ("Learn Docker", [
            "Topics: Containers, Images, Dockerfiles, Volumes, Networks.",
            "Resources:",
            "Docker Official Docs: https://docs.docker.com/get-started/",
            "YouTube: \"Docker Tutorial for Beginners by TechWorld with Nana\" (https://www.youtube.com/watch?v=3c-iBn73dDE)"
        ]),
        ("Learn Kubernetes", [
            "Topics: Pods, Deployments, Services, Ingress, ConfigMaps, Helm.",
            "Resources:",
            "Kubernetes Official Docs: https://kubernetes.io/docs/home/",
            "YouTube: \"Kubernetes for Absolute Beginners by Mumshad Mannambeth\" (https://www.youtube.com/watch?v=s_o8dwzRlu4)"
        ])
    ]),
    ("Infrastructure as Code (IaC)", [
        ("Learn Terraform", [
            "Topics: Providers, Modules, State Management.",
            "Resources:",
            "Terraform Official Docs: https://developer.hashicorp.com/terraform/docs",
            "YouTube: \"Terraform Tutorial for Beginners by FreeCodeCamp\" (https://www.youtube.com/watch?v=SLB_c_ayRMo)"
        ]),
        ("Learn Ansible", [
            "Topics: Playbooks, Inventory, Modules, Roles.",
            "Resources:",
            "Ansible Official Docs: https://docs.ansible.com/",
            "YouTube: \"Ansible Tutorial for Beginners by Simplilearn\" (https://www.youtube.com/watch?v=wgQ3rH9vM9c)"
        ])
    ]),
    ("Cloud Platforms", [
        ("Learn Basics of Cloud (AWS/GCP/Azure)", [
            "Topics: EC2, S3, IAM, RDS, Networking.",
            "Resources:",
            "AWS: \"AWS Certified Cloud Practitioner Full Course by FreeCodeCamp\" (https://www.youtube.com/watch?v=3hLmDS179YE)",
            "GCP: \"Google Cloud Platform Tutorial by Simplilearn\" (https://www.youtube.com/watch?v=lSXuMCB1pTE)"
        ])
    ]),
    ("Monitoring and Logging", [
        ("Learn Prometheus & Grafana", [
            "Topics: Metrics, Dashboards, Alerts.",
            "Resources:",
            "Prometheus Official Docs: https://prometheus.io/docs/",
            "YouTube: \"Prometheus and Grafana Tutorial by TechWorld with Nana\" (https://www.youtube.com/watch?v=RRy3EHzPBgE)"
        ]),
        ("Learn ELK Stack (Elasticsearch, Logstash, Kibana)", [
            "Topics: Log Collection, Parsing, Visualization.",
            "Resources:",
            "Elastic Official Docs: https://www.elastic.co/guide/index.html",
            "YouTube: \"ELK Stack Tutorial by Edureka\" (https://www.youtube.com/watch?v=41DLsB1rpUk)"
        ])
    ]),
    ("Scripting & Automation", [
        ("Learn Bash/Python", [
            "Topics: Automation Scripts, API Calls, File Handling.",
            "Resources:",
            "YouTube: \"Shell Scripting Crash Course by Traversy Media\" (https://www.youtube.com/watch?v=hwrnmQumtPw)",
            "YouTube: \"Python for DevOps by Edureka\" (https://www.youtube.com/watch?v=5eXq-UNKucc)"
        ])
    ]),
    ("Advanced Topics", [
        ("Learn Security in DevOps", [
            "Topics: Secrets Management, SSL/TLS, Vulnerability Scanning.",
            "Resources:",
            "YouTube: \"DevSecOps Tutorial by Simplilearn\" (https://www.youtube.com/watch?v=SC5Zy5aogc4)"
        ]),
        ("Learn Advanced CI/CD", [
            "Tools: CircleCI, GitLab CI, ArgoCD.",
            "Resources:",
            "CircleCI Docs: https://circleci.com/docs/",
            "YouTube: \"GitLab CI/CD Tutorial by Amigoscode\" (https://www.youtube.com/watch?v=PGyhBwLyK2U)"
        ])
    ]),
    ("Practice & Certification", [
        ("Practice Platforms", [
            "KodeKloud: Hands-on Labs for Docker, Kubernetes, etc. (https://kodekloud.com/)",
            "AWS Free Tier: Practice cloud deployments (https://aws.amazon.com/free/)",
            "Katacoda: Interactive learning scenarios (https://www.katacoda.com/)"
        ]),
        ("Certifications", [
            "Docker Certified Associate",
            "AWS Certified Solutions Architect",
            "Certified Kubernetes Administrator (CKA)",
            "HashiCorp Certified: Terraform Associate"
        ])
    ])
]

# Populate the document with sections and subsections
for section_title, subsections in sections:
    doc.add_heading(section_title, level=2)
    for subsection_title, points in subsections:
        doc.add_heading(subsection_title, level=3)
        for point in points:
            doc.add_paragraph(point, style='List Bullet')

# Save the document to a file
file_path = "C:\\Users\\hp\\Desktop\\Beginner_to_Advanced_DevOps_Roadmap.docx"
doc.save(file_path)
file_path
